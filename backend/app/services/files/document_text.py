"""Extract plain text from binary document formats for metadata and scanning."""

import io

from fastapi import HTTPException, status
from pypdf import PdfReader
from docx import Document


def extract_pdf_text(content: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or unreadable PDF file",
        ) from exc

    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts), len(reader.pages)


def extract_docx_text(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or unreadable Word document",
        ) from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        if stripped:
            parts.append(stripped)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped:
                    parts.append(stripped)
    return "\n".join(parts)


def document_lines(content: bytes, file_type: str) -> tuple[list[str], dict]:
    extra: dict = {"format": file_type}
    if file_type == "pdf":
        text, page_count = extract_pdf_text(content)
        extra["page_count"] = page_count
    elif file_type == "docx":
        text = extract_docx_text(content)
        extra["paragraph_count"] = len([ln for ln in text.splitlines() if ln.strip()])
    else:
        raise ValueError(f"Unsupported document type: {file_type}")

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    extra["line_count"] = len(lines)
    extra["char_count"] = len(text)
    return lines, extra
