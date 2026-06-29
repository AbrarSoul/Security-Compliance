"""Extract PDF and DOCX documents with structure preservation."""

from __future__ import annotations

import io

import pdfplumber
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.files.extraction.types import ExtractedContent
from app.services.files.extraction.utils import normalize_cell


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    return "heading" in style_name or style_name.startswith("title")


def extract_pdf(content: bytes) -> ExtractedContent:
    text_blocks: list[dict] = []
    tables: list[dict] = []
    pages_meta: list[dict] = []
    headings: list[dict] = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            pages_meta.append({"page": page_num, "char_count": len(page_text)})

            for line_idx, line in enumerate(page_text.splitlines(), start=1):
                stripped = line.strip()
                if not stripped:
                    continue
                block = {"text": stripped, "page": page_num, "line": line_idx}
                if len(stripped) < 120 and stripped == stripped.upper() and len(stripped.split()) <= 12:
                    headings.append({"text": stripped, "page": page_num, "line": line_idx})
                    block["is_heading"] = True
                text_blocks.append(block)

            for table_idx, table in enumerate(page.extract_tables() or [], start=1):
                rows = [[normalize_cell(cell) for cell in row] for row in table if row]
                if not rows:
                    continue
                headers = rows[0]
                tables.append({
                    "name": f"page_{page_num}_table_{table_idx}",
                    "page": page_num,
                    "table_index": table_idx,
                    "headers": headers,
                    "rows": rows[1:51],
                    "row_count": max(len(rows) - 1, 0),
                })

    sections = [h["text"] for h in headings[:50]]
    return ExtractedContent(
        file_type="pdf",
        text_blocks=text_blocks,
        tables=tables,
        metadata={"page_count": len(pages_meta), "table_count": len(tables)},
        structure={"headings": headings, "sections": sections, "pages": pages_meta, "sheets": []},
    )


def _paragraph_heading_level(paragraph: Paragraph) -> int | None:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "heading" in style_name:
        for level in range(1, 7):
            if str(level) in style_name:
                return level
        return 1
    if style_name.startswith("title"):
        return 1
    return None


def extract_docx(content: bytes) -> ExtractedContent:
    document = Document(io.BytesIO(content))
    text_blocks: list[dict] = []
    tables: list[dict] = []
    headings: list[dict] = []
    paragraph_index = 0

    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        if not stripped:
            continue
        paragraph_index += 1
        block: dict = {"text": stripped, "paragraph_index": paragraph_index}
        level = _paragraph_heading_level(paragraph)
        if level is not None:
            block["is_heading"] = True
            block["heading_level"] = level
            headings.append({
                "text": stripped,
                "level": level,
                "paragraph_index": paragraph_index,
            })
        text_blocks.append(block)

    for table_index, table in enumerate(document.tables, start=1):
        table_dict = _extract_docx_table(table, table_index)
        tables.append(table_dict)

    sections = [h["text"] for h in headings]
    return ExtractedContent(
        file_type="docx",
        text_blocks=text_blocks,
        tables=tables,
        metadata={
            "paragraph_count": paragraph_index,
            "table_count": len(tables),
            "heading_count": len(headings),
        },
        structure={"headings": headings, "sections": sections, "pages": [], "sheets": []},
    )


def _extract_docx_table(table: Table, table_index: int) -> dict:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([normalize_cell(cell.text) for cell in row.cells])
    headers = rows[0] if rows else []
    return {
        "name": f"table_{table_index}",
        "table_index": table_index,
        "headers": headers,
        "rows": rows[1:51],
        "row_count": max(len(rows) - 1, 0),
    }
