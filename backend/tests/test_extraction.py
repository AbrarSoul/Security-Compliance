"""Unit tests for file extraction across all supported types."""

from __future__ import annotations

import io
import json

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.services.files.extraction import (
    extract_csv,
    extract_docx,
    extract_json,
    extract_jsonl,
    extract_log,
    extract_markdown,
    extract_ndjson,
    extract_pdf,
    extract_tsv,
    extract_txt,
    extract_xml,
    extract_xlsx,
    extract_yaml,
    extract_yml,
    extract_file,
)
from app.services.files.excel_io import build_sample_xlsx


def _normalized_keys(result) -> set[str]:
    return {"file_type", "text_blocks", "tables", "metadata", "records", "structure"}


def _assert_normalized(result, file_type: str) -> None:
    assert result.file_type == file_type
    data = result.to_dict()
    assert _normalized_keys(result) <= set(data.keys())
    assert isinstance(result.text_blocks, list)
    assert isinstance(result.tables, list)
    assert isinstance(result.metadata, dict)
    assert isinstance(result.records, list)
    assert isinstance(result.structure, dict)


def test_extract_csv():
    content = b"name,email\nAlice,alice@example.com\nBob,bob@example.com\n"
    result = extract_csv(content)
    _assert_normalized(result, "csv")
    assert result.data_quality is not None
    assert result.data_quality["row_count"] == 2
    assert result.data_quality["column_count"] == 2
    assert "name" in result.structure["headings"]


def test_extract_tsv():
    content = b"name\temail\nAlice\talice@example.com\n"
    result = extract_tsv(content)
    _assert_normalized(result, "tsv")
    assert result.data_quality["row_count"] == 1


def test_extract_json_array():
    content = json.dumps([{"name": "Alice"}, {"name": "Bob"}]).encode()
    result = extract_json(content)
    _assert_normalized(result, "json")
    assert len(result.records) == 2


def test_extract_json_object():
    content = json.dumps({"policy": "retention", "version": 1}).encode()
    result = extract_json(content)
    assert len(result.records) == 1
    assert "policy" in result.records[0]


def test_extract_jsonl():
    content = b'{"email": "a@example.com"}\n{"email": "b@example.com"}\n'
    result = extract_jsonl(content)
    _assert_normalized(result, "jsonl")
    assert len(result.records) == 2


def test_extract_ndjson():
    content = b'{"id": 1}\n{"id": 2}\n'
    result = extract_ndjson(content)
    _assert_normalized(result, "ndjson")
    assert result.file_type == "ndjson"


def test_extract_txt():
    content = b"Line one\nLine two\n"
    result = extract_txt(content)
    _assert_normalized(result, "txt")
    assert len(result.text_blocks) == 2


def test_extract_xml():
    content = b"""<records>
      <item><name>Alice</name><email>a@example.com</email></item>
      <item><name>Bob</name><email>b@example.com</email></item>
    </records>"""
    result = extract_xml(content)
    _assert_normalized(result, "xml")
    assert len(result.records) == 2


def test_extract_yaml():
    content = b"- name: Alice\n  email: alice@example.com\n"
    result = extract_yaml(content)
    _assert_normalized(result, "yaml")
    assert len(result.records) == 1


def test_extract_yml():
    content = b"name: Alice\nemail: alice@example.com\n"
    result = extract_yml(content)
    _assert_normalized(result, "yml")
    assert len(result.records) == 1


def test_extract_log():
    content = (
        b"2024-01-01 10:00:00 INFO user login\n"
        b"2024-01-01 10:01:00 ERROR authentication failed\n"
        b"2024-01-01 10:02:00 WARNING slow response\n"
    )
    result = extract_log(content)
    _assert_normalized(result, "log")
    assert result.metadata["level_counts"]["ERROR"] == 1
    assert any(f["type"] == "authentication_failure" for f in result.metadata["log_findings"])


def test_extract_markdown():
    content = b"# Title\n\n- bullet\n\n```python\nprint('hi')\n```\n\n[link](https://example.com)\n"
    result = extract_markdown(content)
    _assert_normalized(result, "md")
    assert result.metadata["heading_count"] == 1
    assert result.metadata["code_block_count"] == 1
    assert result.metadata["link_count"] == 1


def _sample_pdf(text: str = "Data retention policy and privacy notice") -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, text)
    pdf.save()
    return buffer.getvalue()


def test_extract_pdf():
    result = extract_pdf(_sample_pdf())
    _assert_normalized(result, "pdf")
    assert len(result.text_blocks) >= 1
    assert result.structure["pages"]


def _sample_docx(text: str = "Access control and encryption policy") -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Security Policy", level=1)
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def test_extract_docx():
    result = extract_docx(_sample_docx())
    _assert_normalized(result, "docx")
    assert len(result.text_blocks) >= 2
    assert result.metadata["heading_count"] >= 1


def test_extract_xlsx():
    content = build_sample_xlsx([["email", "name"], ["alice@example.com", "Alice"]])
    result = extract_xlsx(content)
    _assert_normalized(result, "xlsx")
    assert result.data_quality["row_count"] == 1
    assert "Sheet1" in result.structure["sheets"]


def test_extract_file_dispatcher():
    content = b"a,b\n1,2\n"
    result = extract_file("csv", content)
    assert result.file_type == "csv"


def test_csv_detects_missing_values():
    content = b"name,email\nAlice,\nBob,bob@example.com\n"
    result = extract_csv(content)
    dq_types = {f["type"] for f in result.data_quality["findings"]}
    assert "missing_values" in dq_types
    email_finding = next(f for f in result.data_quality["findings"] if f["column"] == "email")
    locations = email_finding["evidence"]["locations"]
    assert locations[0]["index"] == 2
    assert locations[0]["column"] == "email"
    assert "empty" in locations[0]["preview"].lower() or locations[0]["value"] == "(empty)"


def test_csv_detects_duplicates():
    content = b"name,email\nAlice,a@x.com\nAlice,a@x.com\n"
    result = extract_csv(content)
    assert result.data_quality["duplicate_rows"] == 1
